import json
import os
import uuid

import fitz
import chardet
from docx import Document as DocxDocument
from fastapi import HTTPException
from sqlalchemy.ext.asyncio import AsyncSession

from src.models.document import Document, DocumentChunk
from src.repositories.document_repo import DocumentChunkRepository, DocumentRepository
from src.workflows.document_pipeline import process_document

import structlog

logger = structlog.get_logger()

UPLOAD_DIR = "uploads"
os.makedirs(UPLOAD_DIR, exist_ok=True)

ALLOWED_FILE_TYPES = {"pdf", "docx", "txt", "md", "csv", "json", "html"}


async def upload_document(session: AsyncSession, user_id: int, filename: str, file_content: bytes) -> Document:
    file_type = filename.rsplit(".", 1)[-1].lower() if "." in filename else "txt"
    if file_type not in ALLOWED_FILE_TYPES:
        raise HTTPException(status_code=400, detail=f"Unsupported file type: {file_type}. Allowed: {', '.join(sorted(ALLOWED_FILE_TYPES))}")
    file_id = uuid.uuid4().hex
    file_path = os.path.join(UPLOAD_DIR, f"{file_id}.{file_type}")

    with open(file_path, "wb") as f:
        f.write(file_content)

    doc_repo = DocumentRepository(session)
    doc = Document(
        user_id=user_id,
        filename=filename,
        file_type=file_type,
        file_path=file_path,
        status="pending",
    )
    doc = await doc_repo.create(doc)  # flush happens here

    logger.info("document_uploaded", doc_id=doc.id, filename=filename)
    return doc


async def parse_document(file_path: str, file_type: str) -> str:
    if file_type == "pdf":
        return _parse_pdf(file_path)
    elif file_type == "docx":
        return _parse_docx(file_path)
    else:
        return _parse_text(file_path)


def _parse_pdf(file_path: str) -> str:
    doc = fitz.open(file_path)
    text_parts = []
    for page in doc:
        text_parts.append(page.get_text())
    doc.close()
    return "\n".join(text_parts)


def _parse_docx(file_path: str) -> str:
    doc = DocxDocument(file_path)
    return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])


def _parse_text(file_path: str) -> str:
    with open(file_path, "rb") as f:
        raw = f.read()
    detected = chardet.detect(raw)
    encoding = detected.get("encoding", "utf-8")
    return raw.decode(encoding)


async def process_document_upload(session: AsyncSession, document_id: int) -> dict:
    doc_repo = DocumentRepository(session)
    doc = await doc_repo.get(document_id)
    if not doc:
        return {"error": "Document not found"}

    try:
        content = await parse_document(doc.file_path, doc.file_type)
        result = await process_document(document_id, content, doc.filename, session=session)
        await doc_repo.update(doc, status="completed", chunk_count=result["chunk_count"])
        await session.commit()
        return result
    except Exception as e:
        await doc_repo.update(doc, status="failed")
        await session.commit()
        logger.error("document_processing_failed", doc_id=document_id, error=str(e))
        return {"error": str(e)}


async def delete_document(session: AsyncSession, document_id: int) -> bool:
    doc_repo = DocumentRepository(session)
    chunk_repo = DocumentChunkRepository(session)
    doc = await doc_repo.get(document_id)
    if not doc:
        return False

    # Get chunk IDs before deleting from DB
    chunks = await chunk_repo.get_document_chunks(document_id)
    chunk_ids = [c.id for c in chunks]

    # Remove from BM25 index
    from src.retrieval.bm25 import get_bm25_index
    bm25 = get_bm25_index()
    bm25.remove_by_chunk_ids(set(chunk_ids))

    # Remove from Milvus
    try:
        from src.retrieval.vector_store import ensure_collection
        collection = ensure_collection()
        collection.delete(f"document_id == {document_id}")
        collection.flush()
    except Exception as e:
        logger.warning("milvus_delete_failed", doc_id=document_id, error=str(e))

    # Remove from PostgreSQL
    await chunk_repo.delete_document_chunks(document_id)
    if os.path.exists(doc.file_path):
        os.remove(doc.file_path)
    await doc_repo.delete(doc)
    await session.commit()
    return True
